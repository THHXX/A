"""Convert MiroFish_需求规格说明书.md to a formal docx with images embedded.

Handles:
- Headings (# .. ######) → Heading 1..6
- Paragraphs with **bold** / *italic* / `code` inline
- Tables (| a | b |)
- Images ![alt](path)
- Mermaid code fences → pre-rendered PNGs from _mermaid_imgs/
- Other code fences → fixed-width preformatted block
- Horizontal rules (---) → page break-ish separator paragraph
- Blockquotes (>) → italic indented paragraph
- Numbered / bulleted lists (1. / -)
"""
from __future__ import annotations

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_DIR = Path(__file__).parent
MD_FILE = DOC_DIR / "毕设论文_MiroFish多智能体仿真预测系统的设计与实现.md"
IMG_DIR = DOC_DIR
MERMAID_DIR = DOC_DIR / "_mermaid_imgs"
CH3_IMG_DIR = DOC_DIR / "图片"
OUT_FILE = DOC_DIR / "毕设论文_MiroFish多智能体仿真预测系统的设计与实现.docx"

THESIS_TITLE = "MiroFish多智能体仿真预测系统的设计与实现"

CN_FONT = "宋体"
EN_FONT = "Times New Roman"
HEADING_CN_FONT = "黑体"
CODE_FONT = "Consolas"


def set_run_fonts(run, *, cn=CN_FONT, en=EN_FONT, size=None, bold=None, italic=None, color=None):
    """Apply Chinese + Western font split + size/style."""
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)
    rFonts.set(qn("w:ascii"), en)
    rFonts.set(qn("w:hAnsi"), en)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_inline_runs(paragraph, text, *, base_size=10.5, code_color=RGBColor(0xC7, 0x25, 0x4E)):
    """Parse **bold**, *italic*, `code` and emit runs."""
    # Tokenize
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_run_fonts(run, size=base_size)
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_fonts(run, size=base_size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_fonts(run, cn=CODE_FONT, en=CODE_FONT, size=base_size - 0.5, color=code_color)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_fonts(run, size=base_size, italic=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_fonts(run, size=base_size)


def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    sizes = {1: 16, 2: 14, 3: 13, 4: 12, 5: 11, 6: 10.5}
    size = sizes.get(level, 10.5)
    for run in h.runs:
        set_run_fonts(run, cn=HEADING_CN_FONT, en=EN_FONT, size=size, bold=True,
                      color=RGBColor(0x00, 0x00, 0x00) if level <= 3 else RGBColor(0x33, 0x33, 0x33))
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraph(doc, text, *, list_kind=None, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if list_kind == "bullet":
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.37)
        run = p.add_run("• ")
        set_run_fonts(run, size=size)
        add_inline_runs(p, text, base_size=size)
    elif list_kind and list_kind.startswith("num:"):
        num = list_kind.split(":", 1)[1]
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        run = p.add_run(f"{num}. ")
        set_run_fonts(run, size=size)
        add_inline_runs(p, text, base_size=size)
    else:
        p.paragraph_format.first_line_indent = Cm(0.74)
        add_inline_runs(p, text, base_size=size)
    return p


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_fonts(run, size=10, italic=True, color=RGBColor(0x55, 0x55, 0x55))


def _make_borderless_table(doc, rows=1, cols=1):
    """Create a borderless table whose first row won't split across pages."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)
    trPr = table.rows[0]._tr.get_or_add_trPr()
    cantSplit = OxmlElement("w:cantSplit")
    cantSplit.set(qn("w:val"), "true")
    trPr.append(cantSplit)
    return table


def add_image(doc, img_path: Path, *, width_cm=14, caption=None):
    table = _make_borderless_table(doc, rows=1, cols=1)
    cell = table.rows[0].cells[0]
    ip = cell.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(6)
    ip.paragraph_format.space_after = Pt(2)
    irun = ip.add_run()
    irun.add_picture(str(img_path), width=Cm(width_cm))
    if caption:
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cap_m = re.match(r"^(图 \d+-\d+)\s*(.*)", caption)
        if cap_m:
            run1 = cp.add_run(cap_m.group(1) + " ")
            set_run_fonts(run1, size=9, bold=True)
            run2 = cp.add_run(cap_m.group(2))
            set_run_fonts(run2, size=9)
        else:
            crun = cp.add_run(caption)
            set_run_fonts(crun, size=9)
    doc.add_paragraph()


def add_ui_design_card(doc, img_path, caption, desc_text):
    """Render a UI design item as two-column table: left=image, right=description."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set column widths: left ~45% (image), right ~55% (text)
    for row in table.rows:
        row.cells[0].width = Cm(7.2)
        row.cells[1].width = Cm(8.3)

    # Left cell: image
    left_cell = table.rows[0].cells[0]
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(4)
    lr = lp.add_run()
    lr.add_picture(str(img_path), width=Cm(6.8))

    # Caption below image
    cp = left_cell.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    cr = cp.add_run(caption)
    set_run_fonts(cr, size=8, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    # Right cell: description text
    right_cell = table.rows[0].cells[1]
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = right_cell.paragraphs[0]
    rp.paragraph_format.line_spacing = 1.35
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(4)
    add_inline_runs(rp, desc_text, base_size=9.5)

    # Borders
    for cell in table.rows[0].cells:
        set_cell_borders(cell)

    doc.add_paragraph()  # spacer


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    text = "\n".join(code_lines)
    run = p.add_run(text)
    set_run_fonts(run, cn=CODE_FONT, en=CODE_FONT, size=9, color=RGBColor(0x33, 0x33, 0x33))


def add_md_table(doc, header_cells, body_rows):
    cols = len(header_cells)
    table = doc.add_table(rows=1 + len(body_rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header
    for i, h in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        run = cell.paragraphs[0].add_run(h.strip())
        set_run_fonts(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, "1F3A68")
        set_cell_borders(cell)

    # Body
    for r, row in enumerate(body_rows, start=1):
        for c, val in enumerate(row[:cols]):
            cell = table.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            # Markdown <br> in cells → paragraph break
            segments = val.split("<br>")
            for k, seg in enumerate(segments):
                if k == 0:
                    add_inline_runs(para, seg.strip(), base_size=10)
                else:
                    np = cell.add_paragraph()
                    np.paragraph_format.space_before = Pt(0)
                    np.paragraph_format.space_after = Pt(0)
                    add_inline_runs(np, seg.strip(), base_size=10)
            if r % 2 == 0:
                set_cell_shading(cell, "F5F8FC")
            set_cell_borders(cell)


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— — —")
    set_run_fonts(run, size=10, color=RGBColor(0xAA, 0xAA, 0xAA))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)


def add_cover(doc, title):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_fonts(run, cn=HEADING_CN_FONT, en=EN_FONT, size=26, bold=True,
                  color=RGBColor(0x1F, 0x3A, 0x68))
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("毕业论文（设计）")
    set_run_fonts(srun, cn=HEADING_CN_FONT, en=EN_FONT, size=16, bold=False,
                  color=RGBColor(0x55, 0x55, 0x55))
    doc.add_page_break()


def add_toc(doc):
    """Insert a Word TOC field (users must right-click → Update Field in Word)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("目  录")
    set_run_fonts(run, cn=HEADING_CN_FONT, en=EN_FONT, size=16, bold=True)
    doc.add_paragraph()

    toc_p = doc.add_paragraph()
    run0 = toc_p.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run0._r.append(fldChar_begin)

    run1 = toc_p.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run1._r.append(instrText)

    run2 = toc_p.add_run()
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    run2._r.append(fldChar_sep)

    run3 = toc_p.add_run("（在 Word 中右键点击此处 → 更新域 以生成目录）")
    set_run_fonts(run3, size=9, color=RGBColor(0x99, 0x99, 0x99))

    run4 = toc_p.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run4._r.append(fldChar_end)

    doc.add_page_break()


# ----------------------- Parser -----------------------

def parse_and_render():
    md_text = MD_FILE.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()
    configure_document(doc)
    add_cover(doc, THESIS_TITLE)
    add_toc(doc)

    i = 0
    mermaid_idx = 0
    plantuml_idx = 0
    in_ui_design = False
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Skip the very first H1 since it's already on the cover
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading = stripped.lstrip("# ").strip()
            if heading == THESIS_TITLE:
                i += 1
                continue
            add_heading(doc, heading, 1)
            i += 1
            continue

        # Headings ## ... ######
        m = re.match(r"^(#{2,6})\s+(.+?)\s*$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            add_heading(doc, heading_text, level)
            # Enter UI design section
            if heading_text == "界面设计类图":
                in_ui_design = True
            # Exit UI design section on next major heading
            elif in_ui_design and level <= 4:
                in_ui_design = False
            i += 1
            continue

        # Horizontal rule — also exits UI design section
        if stripped.strip() == "---":
            add_separator(doc)
            in_ui_design = False
            i += 1
            continue

        # UI design two-column card: **（N）UIxx ...**
        ui_item_m = re.match(r"^\*\*（(\d+)）(UI\d{2}\s+.+?)\*\*$", stripped)
        if in_ui_design and ui_item_m:
            j = i + 1
            img_src = None
            caption = ""
            desc_lines = []
            in_right_col = False  # True once we enter the right-side <div>
            while j < len(lines):
                s = lines[j].rstrip()
                # Skip HTML tags
                if re.match(r"^\s*</?div[^>]*>\s*$", s):
                    # If we were collecting description and hit </div>, we're done
                    if in_right_col and s.strip().startswith("</div"):
                        j += 1
                        break
                    # Opening the right-col div (flex: 1)
                    if 'flex: 1' in s or 'flex:1' in s:
                        in_right_col = True
                    j += 1
                    continue
                # Image
                if not img_src and s.startswith("![") and "(" in s:
                    m2 = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", s)
                    if m2:
                        img_src = m2.group(2)
                    j += 1
                    continue
                # Caption
                if not caption and s.startswith("**图 ") and s.endswith("**"):
                    caption = s.strip("*")
                    j += 1
                    continue
                # Stop on next UI card title
                if desc_lines and re.match(r"^\*\*（\d+）UI", s):
                    break
                # Stop on next section heading
                if desc_lines and re.match(r"^#{1,6}\s", s):
                    break
                # Collect description lines in the right column
                if in_right_col and s.strip():
                    desc_lines.append(s.strip())
                j += 1

            # Also collect 操作方法：line after the card
            ops_line = ""
            if j < len(lines) and lines[j].strip().startswith("操作方法："):
                ops_line = lines[j].strip()
                j += 1

            if img_src and desc_lines:
                img_path = (IMG_DIR / img_src).resolve()
                if img_path.exists():
                    desc_text = "；".join(desc_lines) + "。"
                    if ops_line:
                        desc_text += " " + ops_line + "。"
                    add_ui_design_card(doc, img_path, caption, desc_text)
            i = j
            continue

        # Code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].rstrip().startswith("```"):
                buf.append(lines[j])
                j += 1
            if lang == "mermaid":
                mermaid_idx += 1
                img = MERMAID_DIR / f"mermaid_{mermaid_idx}.png"
                if img.exists():
                    add_image(doc, img, width_cm=15.5, caption=f"图：流程图 {mermaid_idx}")
                else:
                    add_code_block(doc, buf)
            elif lang == "plantuml":
                plantuml_idx += 1
                if plantuml_idx <= 10:
                    img = MERMAID_DIR / f"{plantuml_idx}.png"
                    caption = f"图：UC-{plantuml_idx:02d} 顺序图"
                else:
                    img = MERMAID_DIR / f"mermaid_{plantuml_idx}.png"
                    caption = f"图：分析类图"
                if img.exists():
                    add_image(doc, img, width_cm=15.5, caption=caption)
                else:
                    add_code_block(doc, buf)
            else:
                add_code_block(doc, buf)
            i = j + 1
            continue

        # Image
        m = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m:
            alt, src = m.group(1), m.group(2)
            img_path = (IMG_DIR / src).resolve()
            if img_path.exists():
                caption = ""
                peek = i + 1
                while peek < len(lines) and lines[peek].strip() == "":
                    peek += 1
                if peek < len(lines):
                    cap_m = re.match(r"^\*\*图 (\d+-\d+)\*\*\s*(.*)", lines[peek].strip())
                    if cap_m:
                        caption = f"图 {cap_m.group(1)} {cap_m.group(2)}".strip()
                        i = peek  # consume caption line
                if not caption:
                    caption = alt if alt and alt != "alt text" else ""
                add_image(doc, img_path, width_cm=14, caption=caption)
            i += 1
            continue

        # Table
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?\s*[:\-\|\s]+\|[:\-\|\s]+\s*$", lines[i + 1]):
            header = [c.strip() for c in stripped.strip().strip("|").split("|")]
            j = i + 2
            rows = []
            while j < len(lines) and "|" in lines[j] and lines[j].strip().startswith("|"):
                cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                rows.append(cells)
                j += 1
            add_md_table(doc, header, rows)
            i = j
            continue

        # Blockquote
        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].rstrip().startswith("> "):
                quote_lines.append(lines[i].rstrip()[2:])
                i += 1
            add_blockquote(doc, " ".join(quote_lines))
            continue

        # Numbered list item
        m = re.match(r"^\s*(\d+)\.\s+(.+)$", stripped)
        if m:
            num, content = m.group(1), m.group(2)
            add_paragraph(doc, content, list_kind=f"num:{num}")
            i += 1
            continue

        # Bullet list
        m = re.match(r"^\s*[-*]\s+(.+)$", stripped)
        if m:
            add_paragraph(doc, m.group(1), list_kind="bullet")
            i += 1
            continue

        # Empty line or HTML tag
        if stripped.strip() == "" or re.match(r"^\s*</?div[^>]*>\s*$", stripped):
            i += 1
            continue

        # Standalone bold figure caption (fallback: normally consumed by image peek-ahead)
        cap_m = re.match(r"^\*\*图 (\d+-\d+)\*\*\s*(.*)", stripped)
        if cap_m:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(8)
            run1 = p.add_run(f"图 {cap_m.group(1)} ")
            set_run_fonts(run1, size=9, bold=True)
            run2 = p.add_run(cap_m.group(2))
            set_run_fonts(run2, size=9)
            i += 1
            continue

        # Plain paragraph
        add_paragraph(doc, stripped.strip())
        i += 1

    doc.save(str(OUT_FILE))
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    parse_and_render()
