from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import os

DOC = 'MiroFish_需求分析文档.docx'
OUT = 'MiroFish_需求分析文档_含图.docx'
VIZ = 'static/image/viz'

# (anchor_text_prefix, image_filename, caption, image_width_cm)
INSERTIONS = [
    # 1.2 项目背景：插在第二段（"传统预测方法..."）后
    ('传统预测方法多以统计模型与专家经验为主',
     'viz-01-emergence-architecture.png',
     '图 1-1  群体涌现预测模型 · 系统架构图', 15.5),
    # 2.2 用户角色：插在 "表 2-1" 之后
    ('表 2-1  系统用户角色',
     'viz-03-user-personas.png',
     '图 2-1  用户画像与典型场景对比图', 15.5),
    # 3.2 用例图：插在已有 "图 3-1" 之后
    ('图 3-1  MiroFish 系统用例图',
     'viz-02-feature-mindmap.png',
     '图 3-2  功能与非功能需求思维导图', 15.5),
    # 六、工作流：插在 "图 6-1" 之后
    ('图 6-1  系统首页与五步工作流序列',
     'viz-04-workflow-sequence.png',
     '图 6-2  五步工作流泳道序列图（带决策与时间约束）', 16.0),
    # 六、工作流末尾：再插数据流图
    ('图 6-2  五步工作流泳道序列图（带决策与时间约束）',
     'viz-05-data-flow.png',
     '图 6-3  项目级隔离数据流图（FR-06 历史项目管理）', 16.0),
]

doc = Document(DOC)

def insert_after(ref_para, image_path, caption, width_cm):
    body = ref_para._p.getparent()
    idx = list(body).index(ref_para._p)

    # 1) image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(image_path, width=Cm(width_cm))
    img_xml = img_para._p
    body.remove(img_xml)

    # 2) caption paragraph
    cap_para = doc.add_paragraph(caption)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap_para.runs:
        r.font.size = None
        r.bold = False
    cap_xml = cap_para._p
    body.remove(cap_xml)

    # insert in order: caption first (after image), so reverse-insert
    body.insert(idx + 1, img_xml)
    body.insert(idx + 2, cap_xml)

for prefix, fname, caption, w in INSERTIONS:
    target = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            target = p
            break
    if target is None:
        print(f'!! anchor not found: {prefix}')
        continue
    img_path = os.path.join(VIZ, fname)
    insert_after(target, img_path, caption, w)
    print(f'OK  {caption}  -> after "{prefix[:30]}..."')

doc.save(OUT)
print('\nSAVED', OUT)
